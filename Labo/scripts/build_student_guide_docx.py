#!/usr/bin/env python3
"""Build the distributable student guide from its Markdown source."""

from __future__ import annotations

import re
from pathlib import Path

from docx import Document
from docx.enum.section import WD_ORIENT
from docx.enum.table import WD_CELL_VERTICAL_ALIGNMENT
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.oxml import OxmlElement
from docx.oxml.ns import qn
from docx.shared import Inches, Pt, RGBColor

ROOT = Path(__file__).resolve().parents[2]
SOURCE = ROOT / "Labo" / "docs" / "Guide_etudiant_labo_squat.md"
OUTPUT = ROOT / "Labo" / "Guide_etudiant_labo_squat.docx"

BLUE = "2E74B5"
DARK_BLUE = "1F4D78"
MUTED = "5B6573"
LIGHT_BLUE = "E8EEF5"
LIGHT_GRAY = "F2F4F7"


def set_cell_shading(cell, fill: str) -> None:
    properties = cell._tc.get_or_add_tcPr()
    shading = properties.find(qn("w:shd"))
    if shading is None:
        shading = OxmlElement("w:shd")
        properties.append(shading)
    shading.set(qn("w:fill"), fill)


def set_cell_margins(cell, *, top: int, start: int, bottom: int, end: int) -> None:
    properties = cell._tc.get_or_add_tcPr()
    margins = properties.first_child_found_in("w:tcMar")
    if margins is None:
        margins = OxmlElement("w:tcMar")
        properties.append(margins)
    for edge, value in (
        ("top", top),
        ("start", start),
        ("bottom", bottom),
        ("end", end),
    ):
        node = margins.find(qn(f"w:{edge}"))
        if node is None:
            node = OxmlElement(f"w:{edge}")
            margins.append(node)
        node.set(qn("w:w"), str(value))
        node.set(qn("w:type"), "dxa")


def set_table_geometry(
    table, widths: list[int], *, total_width: int = 9360, indent: int = 120
) -> None:
    table.autofit = False
    properties = table._tbl.tblPr
    width = properties.first_child_found_in("w:tblW")
    if width is None:
        width = OxmlElement("w:tblW")
        properties.append(width)
    width.set(qn("w:w"), str(total_width))
    width.set(qn("w:type"), "dxa")

    table_indent = properties.first_child_found_in("w:tblInd")
    if table_indent is None:
        table_indent = OxmlElement("w:tblInd")
        properties.append(table_indent)
    table_indent.set(qn("w:w"), str(indent))
    table_indent.set(qn("w:type"), "dxa")

    grid = table._tbl.tblGrid
    for child in list(grid):
        grid.remove(child)
    for value in widths:
        column = OxmlElement("w:gridCol")
        column.set(qn("w:w"), str(value))
        grid.append(column)

    for row in table.rows:
        for index, cell in enumerate(row.cells):
            value = widths[index]
            cell.width = Inches(value / 1440)
            properties = cell._tc.get_or_add_tcPr()
            cell_width = properties.first_child_found_in("w:tcW")
            if cell_width is None:
                cell_width = OxmlElement("w:tcW")
                properties.append(cell_width)
            cell_width.set(qn("w:w"), str(value))
            cell_width.set(qn("w:type"), "dxa")
            set_cell_margins(cell, top=80, start=120, bottom=80, end=120)


def add_numbering(document: Document, *, ordered: bool) -> int:
    numbering = document.part.numbering_part.element
    abstract_ids = [
        int(node.get(qn("w:abstractNumId")))
        for node in numbering.findall(qn("w:abstractNum"))
    ]
    num_ids = [int(node.get(qn("w:numId"))) for node in numbering.findall(qn("w:num"))]
    abstract_id = max(abstract_ids, default=-1) + 1
    num_id = max(num_ids, default=0) + 1

    abstract = OxmlElement("w:abstractNum")
    abstract.set(qn("w:abstractNumId"), str(abstract_id))
    multi = OxmlElement("w:multiLevelType")
    multi.set(qn("w:val"), "singleLevel")
    abstract.append(multi)
    level = OxmlElement("w:lvl")
    level.set(qn("w:ilvl"), "0")
    start = OxmlElement("w:start")
    start.set(qn("w:val"), "1")
    level.append(start)
    number_format = OxmlElement("w:numFmt")
    number_format.set(qn("w:val"), "decimal" if ordered else "bullet")
    level.append(number_format)
    level_text = OxmlElement("w:lvlText")
    level_text.set(qn("w:val"), "%1." if ordered else "•")
    level.append(level_text)
    justification = OxmlElement("w:lvlJc")
    justification.set(qn("w:val"), "left")
    level.append(justification)
    paragraph_properties = OxmlElement("w:pPr")
    tabs = OxmlElement("w:tabs")
    tab = OxmlElement("w:tab")
    tab.set(qn("w:val"), "num")
    tab.set(qn("w:pos"), "540")
    tabs.append(tab)
    paragraph_properties.append(tabs)
    indent = OxmlElement("w:ind")
    indent.set(qn("w:left"), "540")
    indent.set(qn("w:hanging"), "270")
    paragraph_properties.append(indent)
    level.append(paragraph_properties)
    abstract.append(level)
    numbering.append(abstract)

    instance = OxmlElement("w:num")
    instance.set(qn("w:numId"), str(num_id))
    abstract_ref = OxmlElement("w:abstractNumId")
    abstract_ref.set(qn("w:val"), str(abstract_id))
    instance.append(abstract_ref)
    numbering.append(instance)
    return num_id


def apply_numbering(paragraph, num_id: int) -> None:
    properties = paragraph._p.get_or_add_pPr()
    numbering = OxmlElement("w:numPr")
    level = OxmlElement("w:ilvl")
    level.set(qn("w:val"), "0")
    number = OxmlElement("w:numId")
    number.set(qn("w:val"), str(num_id))
    numbering.append(level)
    numbering.append(number)
    properties.append(numbering)


def add_inline(paragraph, text: str) -> None:
    parts = re.split(r"(`[^`]+`|\*\*[^*]+\*\*)", text)
    for part in parts:
        if not part:
            continue
        if part.startswith("`") and part.endswith("`"):
            run = paragraph.add_run(part[1:-1])
            run.font.name = "Aptos Mono"
            run.font.size = Pt(9.5)
            run.font.color.rgb = RGBColor.from_string(DARK_BLUE)
        elif part.startswith("**") and part.endswith("**"):
            run = paragraph.add_run(part[2:-2])
            run.bold = True
        else:
            paragraph.add_run(part)


def configure_document(document: Document) -> None:
    section = document.sections[0]
    section.orientation = WD_ORIENT.PORTRAIT
    section.page_width = Inches(8.5)
    section.page_height = Inches(11)
    section.top_margin = Inches(1)
    section.bottom_margin = Inches(1)
    section.left_margin = Inches(1)
    section.right_margin = Inches(1)
    section.header_distance = Inches(0.492)
    section.footer_distance = Inches(0.492)
    section.different_first_page_header_footer = False
    document.settings.odd_and_even_pages_header_footer = False

    normal = document.styles["Normal"]
    normal.font.name = "Calibri"
    normal.font.size = Pt(11)
    normal.paragraph_format.space_after = Pt(6)
    normal.paragraph_format.line_spacing = 1.25

    settings = {
        "Title": (24, BLUE, 0, 8),
        "Heading 1": (16, BLUE, 18, 10),
        "Heading 2": (13, BLUE, 14, 7),
        "Heading 3": (12, DARK_BLUE, 10, 5),
    }
    for name, (size, color, before, after) in settings.items():
        style = document.styles[name]
        style.font.name = "Calibri"
        style.font.size = Pt(size)
        style.font.color.rgb = RGBColor.from_string(color)
        style.font.bold = True
        style.paragraph_format.space_before = Pt(before)
        style.paragraph_format.space_after = Pt(after)
        style.paragraph_format.keep_with_next = True

    for header_part in (
        section.header,
        section.even_page_header,
        section.first_page_header,
    ):
        header = header_part.paragraphs[0]
        header.text = "LABORATOIRE DE BIOMÉCANIQUE — SQUAT 2D"
        header.alignment = WD_ALIGN_PARAGRAPH.RIGHT
        header_run = header.runs[0]
        header_run.font.name = "Calibri"
        header_run.font.size = Pt(8)
        header_run.font.color.rgb = RGBColor.from_string(MUTED)

    for footer_part in (
        section.footer,
        section.even_page_footer,
        section.first_page_footer,
    ):
        footer = footer_part.paragraphs[0]
        footer.clear()
        footer.alignment = WD_ALIGN_PARAGRAPH.RIGHT
        footer.add_run("Guide étudiant  •  ")
        field = OxmlElement("w:fldSimple")
        field.set(qn("w:instr"), "PAGE")
        footer._p.append(field)
        for run in footer.runs:
            run.font.name = "Calibri"
            run.font.size = Pt(8)
            run.font.color.rgb = RGBColor.from_string(MUTED)


def add_code_block(document: Document, lines: list[str]) -> None:
    paragraph = document.add_paragraph()
    paragraph.paragraph_format.space_before = Pt(3)
    paragraph.paragraph_format.space_after = Pt(8)
    paragraph.paragraph_format.left_indent = Inches(0.18)
    paragraph.paragraph_format.right_indent = Inches(0.18)
    properties = paragraph._p.get_or_add_pPr()
    shading = OxmlElement("w:shd")
    shading.set(qn("w:fill"), LIGHT_GRAY)
    properties.append(shading)
    run = paragraph.add_run("\n".join(lines))
    run.font.name = "Aptos Mono"
    run.font.size = Pt(8.5)
    run.font.color.rgb = RGBColor.from_string("273142")


def add_markdown_table(document: Document, rows: list[list[str]]) -> None:
    table = document.add_table(rows=len(rows), cols=len(rows[0]))
    table.style = "Table Grid"
    widths = (
        [7600, 1760] if len(rows[0]) == 2 else [9360 // len(rows[0])] * len(rows[0])
    )
    set_table_geometry(table, widths)
    for row_index, values in enumerate(rows):
        row = table.rows[row_index]
        row._tr.get_or_add_trPr().append(OxmlElement("w:cantSplit"))
        for column_index, value in enumerate(values):
            cell = row.cells[column_index]
            cell.vertical_alignment = WD_CELL_VERTICAL_ALIGNMENT.CENTER
            paragraph = cell.paragraphs[0]
            paragraph.paragraph_format.space_after = Pt(0)
            add_inline(paragraph, value.strip())
            if row_index == 0:
                set_cell_shading(cell, LIGHT_BLUE)
                for run in paragraph.runs:
                    run.bold = True
                    run.font.color.rgb = RGBColor.from_string(DARK_BLUE)
            if column_index == len(values) - 1:
                paragraph.alignment = WD_ALIGN_PARAGRAPH.RIGHT
    document.add_paragraph().paragraph_format.space_after = Pt(0)


def build() -> None:
    lines = SOURCE.read_text(encoding="utf-8").splitlines()
    document = Document()
    configure_document(document)

    index = 0
    active_list_kind: str | None = None
    active_num_id: int | None = None
    active_list_count = 0
    while index < len(lines):
        line = lines[index]
        stripped = line.strip()
        if not stripped:
            active_list_kind = None
            active_num_id = None
            active_list_count = 0
            index += 1
            continue

        if stripped.startswith("```"):
            code: list[str] = []
            index += 1
            while index < len(lines) and not lines[index].strip().startswith("```"):
                code.append(lines[index])
                index += 1
            add_code_block(document, code)
            index += 1
            active_list_kind = None
            active_num_id = None
            active_list_count = 0
            continue

        if (
            stripped.startswith("|")
            and index + 1 < len(lines)
            and re.match(r"^\|[\s:|-]+\|$", lines[index + 1].strip())
        ):
            rows = [[cell.strip() for cell in stripped.strip("|").split("|")]]
            index += 2
            while index < len(lines) and lines[index].strip().startswith("|"):
                rows.append(
                    [
                        cell.strip()
                        for cell in lines[index].strip().strip("|").split("|")
                    ]
                )
                index += 1
            add_markdown_table(document, rows)
            active_list_kind = None
            active_num_id = None
            active_list_count = 0
            continue

        heading = re.match(r"^(#{1,3})\s+(.+)$", stripped)
        if heading:
            level = len(heading.group(1))
            text = heading.group(2)
            if level == 1:
                paragraph = document.add_paragraph(style="Title")
            else:
                paragraph = document.add_heading(level=level - 1)
            add_inline(paragraph, text)
            active_list_kind = None
            active_num_id = None
            active_list_count = 0
            index += 1
            continue

        ordered = re.match(r"^\d+\.\s+(.+)$", stripped)
        bullet = re.match(r"^-\s+(.+)$", stripped)
        if ordered or bullet:
            kind = "ordered" if ordered else "bullet"
            if kind != active_list_kind:
                active_num_id = add_numbering(document, ordered=kind == "ordered")
                active_list_kind = kind
                active_list_count = 0
            paragraph = document.add_paragraph()
            paragraph.paragraph_format.space_after = Pt(4)
            paragraph.paragraph_format.line_spacing = 1.25
            if active_list_count == 0:
                paragraph.paragraph_format.keep_with_next = True
            apply_numbering(paragraph, active_num_id)
            add_inline(paragraph, (ordered or bullet).group(1))
            active_list_count += 1
            index += 1
            continue

        paragraph_lines = [stripped]
        index += 1
        while index < len(lines):
            candidate = lines[index].strip()
            if (
                not candidate
                or candidate.startswith(("#", "```", "|", "- "))
                or re.match(r"^\d+\.\s+", candidate)
            ):
                break
            paragraph_lines.append(candidate)
            index += 1
        paragraph = document.add_paragraph()
        add_inline(paragraph, " ".join(paragraph_lines))
        active_list_kind = None
        active_num_id = None
        active_list_count = 0

    document.core_properties.title = (
        "Laboratoire — Squat 2D, dynamique inverse et interprétation pratique"
    )
    document.core_properties.subject = "Guide étudiant Squat_GUI"
    document.core_properties.author = "Laboratoire de biomécanique"
    document.save(OUTPUT)
    print(OUTPUT)


if __name__ == "__main__":
    build()
